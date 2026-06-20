from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


class ReportDocumentExporter:
    def export_docx(
        self,
        *,
        title: str,
        markdown_text: str,
        destination: str | Path,
    ) -> Path:
        output_path = Path(destination).resolve()
        output_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        document.add_heading(title, level=0)
        for raw_line in markdown_text.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith("### "):
                document.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                document.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                document.add_heading(line[2:], level=1)
            elif line.startswith(("- ", "* ")):
                document.add_paragraph(line[2:], style="List Bullet")
            else:
                document.add_paragraph(line)

        styles = document.styles
        styles["Normal"].font.name = "Microsoft YaHei"
        styles["Normal"].font.size = Pt(10.5)
        document.save(output_path)
        return output_path
